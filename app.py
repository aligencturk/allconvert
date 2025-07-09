"""
Bu Flask uygulaması, çeşitli dosya formatlarını birbirine dönüştürmek için bir web arayüzü sağlar.
Kullanıcılar dosya yükleyebilir, hedef formatı seçebilir ve dönüştürülen dosyayı indirebilir.
"""
import os
from datetime import datetime, timedelta
from flask import (Flask, request, render_template, send_file, flash, redirect, url_for, jsonify)
from dotenv import load_dotenv
import logging
from werkzeug.utils import secure_filename
import platform  # İşletim sistemini kontrol etmek için
import zipfile # PDF'ten JPG'e dönüştürme için
import threading
import time
import re
import requests
from bs4 import BeautifulSoup
import yt_dlp
import json
import imageio_ffmpeg
from pydub import AudioSegment
import subprocess
import shutil
import atexit
from concurrent.futures import ThreadPoolExecutor
import psutil  # Sistem kaynaklarını takip etmek için

# FFmpeg'in yolunu bul ve pydub için ayarla
try:
    ffmpeg_path = imageio_ffmpeg.get_ffmpeg_exe()
    AudioSegment.converter = ffmpeg_path
    logging.info(f"FFmpeg pydub için ayarlandı: {ffmpeg_path}")
except Exception as e:
    ffmpeg_path = None # Eğer bulunamazsa None olarak ayarla
    logging.warning(f"imageio-ffmpeg aracılığıyla FFmpeg bulunamadı: {e}. Ses/video dönüştürme işlemleri başarısız olabilir.")

# .env dosyasındaki ortam değişkenlerini yükle
load_dotenv()

# Uygulama yapılandırması
app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'default-secret-key-for-development')
app.config['DOWNLOAD_FOLDER'] = 'downloads'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50 MB max dosya boyutu
app.config['CLEANUP_INTERVAL_HOURS'] = int(os.getenv('CLEANUP_INTERVAL_HOURS', '1'))  # 1 saat
app.config['FILE_RETENTION_HOURS'] = int(os.getenv('FILE_RETENTION_HOURS', '24'))  # 24 saat
app.config['MAX_CONCURRENT_DOWNLOADS'] = int(os.getenv('MAX_CONCURRENT_DOWNLOADS', '5'))  # Max 5 eşzamanlı indirme
app.config['DISK_USAGE_WARNING_PERCENT'] = int(os.getenv('DISK_USAGE_WARNING_PERCENT', '85'))  # %85 disk uyarısı
app.config['DISK_USAGE_CRITICAL_PERCENT'] = int(os.getenv('DISK_USAGE_CRITICAL_PERCENT', '95'))  # %95 disk kritiği

# Rate limiting için Flask-Limiter
try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address
    
    limiter = Limiter(
        key_func=get_remote_address,
        default_limits=["10 per minute", "100 per hour"],
        storage_uri="memory://"
    )
    limiter.init_app(app)
    logging.info("Rate limiting aktif edildi")
except ImportError:
    logging.warning("Flask-Limiter bulunamadı. Rate limiting devre dışı.")
    limiter = None

# Temel loglama yapılandırması
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Gerekli klasörlerin var olduğundan emin ol
if not os.path.exists(app.config['DOWNLOAD_FOLDER']):
    os.makedirs(app.config['DOWNLOAD_FOLDER'])

# Thread pool executor - eşzamanlı işlemleri sınırlamak için
executor = ThreadPoolExecutor(max_workers=app.config['MAX_CONCURRENT_DOWNLOADS'])

# Session yönetimi - TTL ile otomatik temizleme
class SessionManager:
    def __init__(self):
        self.sessions = {}
        self.session_timestamps = {}
        self.max_session_age = timedelta(hours=2)  # 2 saat
    
    def create_session(self, session_id, data):
        """Yeni session oluştur"""
        self.cleanup_expired_sessions()
        self.sessions[session_id] = data
        self.session_timestamps[session_id] = datetime.now()
    
    def get_session(self, session_id):
        """Session'ı al, yoksa None döndür"""
        self.cleanup_expired_sessions()
        return self.sessions.get(session_id)
    
    def update_session(self, session_id, data):
        """Session'ı güncelle"""
        if session_id in self.sessions:
            self.sessions[session_id].update(data)
            self.session_timestamps[session_id] = datetime.now()
    
    def delete_session(self, session_id):
        """Session'ı sil"""
        self.sessions.pop(session_id, None)
        self.session_timestamps.pop(session_id, None)
    
    def cleanup_expired_sessions(self):
        """Süresi dolmuş session'ları temizle"""
        current_time = datetime.now()
        expired_sessions = [
            sid for sid, timestamp in self.session_timestamps.items()
            if current_time - timestamp > self.max_session_age
        ]
        for sid in expired_sessions:
            self.delete_session(sid)
            logging.info(f"Süresi dolmuş session silindi: {sid}")

# Global session manager
session_manager = SessionManager()

# --- SİSTEM YÖNETİMİ FONKSİYONLARI ---

def check_disk_space():
    """Disk kullanımını kontrol et"""
    try:
        disk_usage = shutil.disk_usage(app.config['DOWNLOAD_FOLDER'])
        used_percent = (disk_usage.used / disk_usage.total) * 100
        
        if used_percent >= app.config['DISK_USAGE_CRITICAL_PERCENT']:
            raise Exception(f"Kritik disk kullanımı: %{used_percent:.1f}! İşlem durduruldu.")
        elif used_percent >= app.config['DISK_USAGE_WARNING_PERCENT']:
            logging.warning(f"Yüksek disk kullanımı: %{used_percent:.1f}")
        
        return used_percent
    except Exception as e:
        logging.error(f"Disk kullanımı kontrol edilemedi: {e}")
        return 0

def cleanup_old_files():
    """Eski dosyaları ve klasörleri temizle"""
    try:
        downloads_dir = app.config['DOWNLOAD_FOLDER']
        cutoff_time = datetime.now() - timedelta(hours=app.config['FILE_RETENTION_HOURS'])
        total_cleaned = 0
        
        for item_name in os.listdir(downloads_dir):
            item_path = os.path.join(downloads_dir, item_name)
            
            if os.path.isdir(item_path):
                # Klasör oluşturma zamanını kontrol et
                try:
                    creation_time = datetime.fromtimestamp(os.path.getctime(item_path))
                    if creation_time < cutoff_time:
                        shutil.rmtree(item_path)
                        total_cleaned += 1
                        logging.info(f"Eski klasör silindi: {item_name}")
                except Exception as e:
                    logging.error(f"Klasör silinemedi {item_name}: {e}")
            
            elif os.path.isfile(item_path) and item_path.endswith('.zip'):
                # ZIP dosyalarını kontrol et
                try:
                    creation_time = datetime.fromtimestamp(os.path.getctime(item_path))
                    if creation_time < cutoff_time:
                        os.remove(item_path)
                        total_cleaned += 1
                        logging.info(f"Eski ZIP dosyası silindi: {item_name}")
                except Exception as e:
                    logging.error(f"ZIP dosyası silinemedi {item_name}: {e}")
        
        if total_cleaned > 0:
            logging.info(f"Temizlik tamamlandı: {total_cleaned} eski öğe silindi")
            
        return total_cleaned
    except Exception as e:
        logging.error(f"Dosya temizleme hatası: {e}")
        return 0

def get_system_stats():
    """Sistem istatistiklerini al"""
    try:
        memory = psutil.virtual_memory()
        disk = shutil.disk_usage(app.config['DOWNLOAD_FOLDER'])
        
        return {
            'memory_percent': memory.percent,
            'disk_percent': (disk.used / disk.total) * 100,
            'active_sessions': len(session_manager.sessions),
            'download_folder_size': sum(
                os.path.getsize(os.path.join(dirpath, filename))
                for dirpath, dirnames, filenames in os.walk(app.config['DOWNLOAD_FOLDER'])
                for filename in filenames
            ) / (1024 * 1024)  # MB cinsinden
        }
    except Exception as e:
        logging.error(f"Sistem istatistikleri alınamadı: {e}")
        return {}

def schedule_cleanup():
    """Periyodik temizleme işlemini başlat"""
    def cleanup_worker():
        while True:
            try:
                # Dosya temizleme
                cleanup_old_files()
                
                # Session temizleme
                session_manager.cleanup_expired_sessions()
                
                # Sistem durumu logla
                stats = get_system_stats()
                if stats:
                    logging.info(f"Sistem durumu - RAM: %{stats.get('memory_percent', 0):.1f}, "
                               f"Disk: %{stats.get('disk_percent', 0):.1f}, "
                               f"Aktif Sessions: {stats.get('active_sessions', 0)}, "
                               f"Downloads Boyutu: {stats.get('download_folder_size', 0):.1f}MB")
                
                # Bir sonraki temizliğe kadar bekle
                time.sleep(app.config['CLEANUP_INTERVAL_HOURS'] * 3600)
                
            except Exception as e:
                logging.error(f"Periyodik temizleme hatası: {e}")
                time.sleep(300)  # Hata durumunda 5 dakika bekle

    # Arka plan thread'ini başlat
    cleanup_thread = threading.Thread(target=cleanup_worker, daemon=True)
    cleanup_thread.start()
    logging.info("Periyodik temizleme sistemi başlatıldı")

# Uygulama başlatıldığında temizleme sistemini başlat
schedule_cleanup()

# Uygulama kapanırken temizlik yap
def cleanup_on_exit():
    """Uygulama kapanırken temizlik yap"""
    logging.info("Uygulama kapanıyor, temizlik yapılıyor...")
    cleanup_old_files()
    executor.shutdown(wait=False)

atexit.register(cleanup_on_exit)

# --- DÖNÜŞTÜRÜCÜ FONKSİYONLARI ---

def convert_word_to_pdf(input_path, output_folder):
    """
    Word belgesini (.docx) PDF formatına dönüştürür.
    Gerekli Kütüphane: pip install docx2pdf
    """
    try:
        from docx2pdf import convert
        output_path = os.path.join(output_folder, os.path.basename(input_path).replace(".docx", ".pdf"))
        logging.info(f"'{input_path}' dosyası '{output_path}' olarak dönüştürülüyor...")
        convert(input_path, output_path)
        if os.path.exists(output_path):
            logging.info("Dönüştürme başarılı.")
            return output_path
        else:
            raise Exception("Dönüştürme sonrası çıktı dosyası bulunamadı.")
    except Exception as e:
        logging.error(f"Word'den PDF'e dönüştürme hatası: {e}")
        # Hatanın daha detaylı izlenmesi için traceback'i de loglayabiliriz.
        import traceback
        logging.error(traceback.format_exc())
        return None

def convert_pdf_to_word(input_path, output_folder):
    """
    PDF dosyasını Word belgesine (.docx) dönüştürür.
    Gerekli Kütüphane: pip install pdf2docx
    """
    try:
        from pdf2docx import Converter
        output_path = os.path.join(output_folder, os.path.basename(input_path).replace(".pdf", ".docx"))
        logging.info(f"'{input_path}' dosyası '{output_path}' olarak dönüştürülüyor...")
        
        # PDF'i dönüştür
        cv = Converter(input_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        
        if os.path.exists(output_path):
            logging.info("Dönüştürme başarılı.")
            return output_path
        else:
            raise Exception("Dönüştürme sonrası çıktı dosyası bulunamadı.")
    except Exception as e:
        logging.error(f"PDF'ten Word'e dönüştürme hatası: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return None

def convert_excel_to_pdf(input_path, output_folder):
    """
    Excel dosyasını (.xlsx) PDF formatına dönüştürür.
    Bu basit dönüştürücü, Excel'in ilk sayfasını bir HTML tablosuna çevirir
    ve bunu PDF yapar. Karmaşık formatlamayı korumayabilir.
    Gerekli Kütüphaneler: pip install pandas openpyxl
    """
    try:
        import pandas as pd
        output_path = os.path.join(output_folder, os.path.basename(input_path).replace(".xlsx", ".pdf"))
        
        # Excel dosyasını oku ve HTML'e dönüştür (daha iyi bir PDF çıktısı için)
        df = pd.read_excel(input_path)
        html_content = df.to_html(index=False, border=1)

        # Geçici bir HTML dosyası oluştur
        temp_html_path = os.path.join(output_folder, "temp.html")
        with open(temp_html_path, "w", encoding="utf-8") as f:
            f.write(html_content)
            
        # Bu kısım için WeasyPrint gibi bir kütüphane gerekir. 
        # Alternatif olarak, Windows'ta COM otomasyonu kullanılabilir.
        # Şimdilik basit bir yer tutucu olarak bırakalım ve COM kullanalım.
        if platform.system() == "Windows":
            import win32com.client
            excel = win32com.client.Dispatch("Excel.Application")
            excel.Visible = False
            
            workbook = excel.Workbooks.Open(os.path.abspath(input_path))
            workbook.ActiveSheet.ExportAsFixedFormat(0, os.path.abspath(output_path))
            workbook.Close(False)
            excel.Quit()
            
            if os.path.exists(output_path):
                logging.info("Excel -> PDF dönüştürme başarılı (COM).")
                return output_path

        raise NotImplementedError("Excel'den PDF'e dönüştürme şu anda sadece Windows'ta ve Microsoft Office yüklü ise desteklenmektedir.")

    except Exception as e:
        logging.error(f"Excel'den PDF'e dönüştürme hatası: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return None

def convert_powerpoint_to_pdf(input_path, output_folder):
    """
    PowerPoint dosyasını (.pptx) PDF formatına dönüştürür.
    Sadece Windows'ta ve Microsoft Office yüklü ise çalışır.
    Gerekli Kütüphane: pip install pypiwin32
    """
    try:
        output_path = os.path.join(output_folder, os.path.basename(input_path).rsplit('.', 1)[0] + ".pdf")
        
        powerpoint = win32com.client.Dispatch("PowerPoint.Application")
        presentation = powerpoint.Presentations.Open(os.path.abspath(input_path), WithWindow=False)
        presentation.SaveAs(os.path.abspath(output_path), 32) # 32 = ppFormatPDF
        presentation.Close()
        powerpoint.Quit()

        if os.path.exists(output_path):
            logging.info("PowerPoint -> PDF dönüştürme başarılı (COM).")
            return output_path
        else:
            raise Exception("Dönüştürme sonrası çıktı dosyası bulunamadı.")
    except Exception as e:
        logging.error(f"PowerPoint'ten PDF'e dönüştürme hatası: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return None

def convert_word_to_txt(input_path, output_folder):
    """
    Word belgesini (.docx) metin dosyasına (.txt) dönüştürür.
    Gerekli Kütüphane: pip install python-docx
    """
    try:
        import docx
        output_path = os.path.join(output_folder, os.path.basename(input_path).replace(".docx", ".txt"))
        doc = docx.Document(input_path)
        full_text = [para.text for para in doc.paragraphs]
        with open(output_path, "w", encoding="utf-8") as f:
            f.write('\n'.join(full_text))
        
        if os.path.exists(output_path):
            logging.info("Word -> TXT dönüştürme başarılı.")
            return output_path
        raise Exception("Dönüştürme sonrası çıktı dosyası bulunamadı.")
    except Exception as e:
        logging.error(f"Word'den TXT'ye dönüştürme hatası: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return None

def convert_txt_to_word(input_path, output_folder):
    """
    Metin dosyasını (.txt) Word belgesine (.docx) dönüştürür.
    Gerekli Kütüphane: pip install python-docx
    """
    try:
        import docx
        output_path = os.path.join(output_folder, os.path.basename(input_path).replace(".txt", ".docx"))
        doc = docx.Document()
        with open(input_path, "r", encoding="utf-8") as f:
            doc.add_paragraph(f.read())
        doc.save(output_path)

        if os.path.exists(output_path):
            logging.info("TXT -> Word dönüştürme başarılı.")
            return output_path
        raise Exception("Dönüştürme sonrası çıktı dosyası bulunamadı.")
    except Exception as e:
        logging.error(f"TXT'den Word'e dönüştürme hatası: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return None

def convert_pdf_to_jpg(input_path, output_folder):
    """
    PDF dosyasının her sayfasını JPG resmine dönüştürür ve bir ZIP dosyası olarak sunar.
    Gerekli Kütüphane: pip install PyMuPDF Pillow
    """
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(input_path)
        image_paths = []
        
        # Eğer hiç sayfa yoksa hata ver
        if len(doc) == 0:
            raise Exception("PDF dosyasında dönüştürülecek sayfa bulunamadı.")

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            pix = page.get_pixmap()
            image_path = os.path.join(output_folder, f"sayfa_{page_num + 1}.jpg")
            pix.save(image_path, "jpeg")
            image_paths.append(image_path)
        doc.close()

        # Tek bir resim varsa doğrudan gönder, birden çoksa ZIP'le
        if len(image_paths) == 1:
            logging.info("PDF -> JPG (tek sayfa) dönüştürme başarılı.")
            return image_paths[0]
        else:
            zip_path = os.path.join(output_folder, os.path.basename(input_path).replace(".pdf", ".zip"))
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for file in image_paths:
                    zipf.write(file, os.path.basename(file))
            
            logging.info(f"PDF -> JPG (çoklu sayfa) dönüştürme başarılı. '{zip_path}' oluşturuldu.")
            return zip_path
            
    except Exception as e:
        logging.error(f"PDF'ten JPG'ye dönüştürme hatası: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return None

def convert_jpg_to_pdf(input_path, output_folder):
    """
    JPG resmini PDF dosyasına dönüştürür.
    Gerekli Kütüphane: pip install Pillow
    """
    try:
        from PIL import Image
        output_path = os.path.join(output_folder, os.path.basename(input_path).replace(".jpg", ".pdf").replace(".jpeg", ".pdf"))
        
        image = Image.open(input_path)
        # RGBA'dan RGB'ye dönüştür, çünkü PDF CMYK veya RGB'yi destekler
        if image.mode == 'RGBA':
            image = image.convert('RGB')
            
        image.save(output_path, "PDF", resolution=100.0)
        
        if os.path.exists(output_path):
            logging.info("JPG -> PDF dönüştürme başarılı.")
            return output_path
        raise Exception("Dönüştürme sonrası çıktı dosyası bulunamadı.")
    except Exception as e:
        logging.error(f"JPG'den PDF'e dönüştürme hatası: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return None

def convert_image_format(input_path, output_folder, target_format):
    """
    Bir resim formatını diğerine dönüştürür (örn: JPG -> PNG).
    Gerekli Kütüphane: pip install Pillow
    """
    try:
        from PIL import Image
        base_name = os.path.basename(input_path).rsplit('.', 1)[0]
        output_path = os.path.join(output_folder, f"{base_name}.{target_format.lower()}")
        
        image = Image.open(input_path)
        # PNG şeffaflığını korumak için RGBA modunu kontrol et
        if image.mode != 'RGB' and target_format.lower() == 'jpg':
             image = image.convert('RGB')
             
        image.save(output_path, format=target_format.upper())
        
        if os.path.exists(output_path):
            logging.info(f"Resim formatı {target_format.upper()} olarak dönüştürüldü.")
            return output_path
        raise Exception("Dönüştürme sonrası çıktı dosyası bulunamadı.")
    except Exception as e:
        logging.error(f"Resim formatı dönüştürme hatası: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return None

# --- Wrapper Fonksiyonlar (CONVERTERS sözlüğü için) ---
def convert_jpg_to_png(input_path, output_folder):
    return convert_image_format(input_path, output_folder, 'PNG')

def convert_png_to_jpg(input_path, output_folder):
    return convert_image_format(input_path, output_folder, 'JPG')

# --- Ses, Video, Veri ve Arşiv Dönüştürücüleri ---

def convert_audio(input_path, output_folder, target_format):
    """
    Ses dosyalarını dönüştürür (örn: MP3 -> WAV).
    Gerekli Kütüphane: pip install pydub
    Gerekli Program: ffmpeg
    """
    try:
        from pydub import AudioSegment
        output_path = os.path.join(output_folder, os.path.basename(input_path).rsplit('.', 1)[0] + f".{target_format}")
        
        logging.info(f"Ses dosyası okunuyor: {input_path}")
        audio = AudioSegment.from_file(input_path)
        
        logging.info(f"Ses dosyası {target_format} formatına dışa aktarılıyor...")
        audio.export(output_path, format=target_format)

        if os.path.exists(output_path):
            logging.info(f"Ses dönüştürme başarılı: -> {output_path}")
            return output_path
        raise Exception("Dönüştürme sonrası çıktı dosyası bulunamadı.")
    except Exception as e:
        logging.error(f"Ses dönüştürme hatası: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return None

def convert_video(input_path, output_folder, target_format):
    """
    Video dosyalarını dönüştürür (örn: MP4 -> AVI).
    Gerekli Kütüphane: pip install moviepy
    Gerekli Program: ffmpeg
    """
    try:
        from moviepy.editor import VideoFileClip
        output_path = os.path.join(output_folder, os.path.basename(input_path).rsplit('.', 1)[0] + f".{target_format}")
        
        logging.info(f"Video dosyası işleniyor: {input_path}")
        with VideoFileClip(input_path) as video:
            video.write_videofile(output_path, codec='libxvid' if target_format == 'avi' else None)

        if os.path.exists(output_path):
            logging.info(f"Video dönüştürme başarılı: -> {output_path}")
            return output_path
        raise Exception("Dönüştürme sonrası çıktı dosyası bulunamadı.")
    except Exception as e:
        logging.error(f"Video dönüştürme hatası: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return None

def convert_json_to_xml(input_path, output_folder):
    """
    JSON dosyasını XML dosyasına dönüştürür.
    Gerekli Kütüphaneler: pip install dicttoxml
    """
    try:
        import json
        from dicttoxml import dicttoxml
        output_path = os.path.join(output_folder, os.path.basename(input_path).replace(".json", ".xml"))

        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        xml_data = dicttoxml(data, custom_root='root', attr_type=False)
        
        with open(output_path, 'wb') as f:
            f.write(xml_data)

        if os.path.exists(output_path):
            logging.info("JSON -> XML dönüştürme başarılı.")
            return output_path
        raise Exception("Dönüştürme sonrası çıktı dosyası bulunamadı.")
    except Exception as e:
        logging.error(f"JSON'dan XML'e dönüştürme hatası: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return None

def convert_xml_to_json(input_path, output_folder):
    """
    XML dosyasını JSON dosyasına dönüştürür.
    Gerekli Kütüphane: pip install xmltodict
    """
    try:
        import json
        import xmltodict
        output_path = os.path.join(output_folder, os.path.basename(input_path).replace(".xml", ".json"))

        with open(input_path, 'r', encoding='utf-8') as f:
            xml_content = f.read()
            
        data_dict = xmltodict.parse(xml_content)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data_dict, f, indent=4)

        if os.path.exists(output_path):
            logging.info("XML -> JSON dönüştürme başarılı.")
            return output_path
        raise Exception("Dönüştürme sonrası çıktı dosyası bulunamadı.")
    except Exception as e:
        logging.error(f"XML'den JSON'a dönüştürme hatası: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return None

def convert_rar_to_zip(input_path, output_folder):
    """
    RAR arşivini ZIP formatına dönüştürür.
    Sistemde 'unrar' komutunun yüklü olmasını gerektirir.
    """
    try:
        # Çıktı için ZIP dosyasının yolu
        output_zip_path = os.path.join(output_folder, os.path.basename(input_path).rsplit('.', 1)[0] + '.zip')
        
        # RAR dosyasını geçici bir klasöre çıkar
        temp_extract_folder = os.path.join(output_folder, 'temp_unrar')
        os.makedirs(temp_extract_folder, exist_ok=True)
        
        # unrar komutunu çalıştır
        # Not: Bu komutun sistemde PATH içinde olması gerekir.
        # Windows için: https://www.rarlab.com/rar/unrarw32.exe
        # Linux için: sudo apt-get install unrar
        cmd = ['unrar', 'x', '-o+', input_path, temp_extract_folder]
        
        # subprocess.run daha modern bir yaklaşımdır
        result = subprocess.run(cmd, check=True, capture_output=True, text=True)
        logging.info(f"unrar çıktısı: {result.stdout}")

        # Çıkarılan dosyaları ZIP'le
        with zipfile.ZipFile(output_zip_path, 'w') as zipf:
            for root, _, files in os.walk(temp_extract_folder):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, temp_extract_folder)
                    zipf.write(file_path, arcname)

        # Geçici klasörü temizle
        shutil.rmtree(temp_extract_folder)

        return output_zip_path
    except FileNotFoundError:
        logging.error("RAR dönüştürme hatası: 'unrar' programı sisteminizde bulunamadı veya PATH içinde değil.")
        raise RuntimeError("Sistemde 'unrar' programı bulunamadığı için RAR dosyaları dönüştürülemiyor.")
    except subprocess.CalledProcessError as e:
        logging.error(f"unrar çalıştırılırken hata: {e.stderr}")
        raise RuntimeError(f"RAR dosyası işlenirken hata oluştu: {e.stderr}")
    except Exception as e:
        logging.error(f"RAR'dan ZIP'e dönüştürme sırasında genel hata: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return None

def convert_wav_to_mp3(input_path, output_folder):
    """WAV'ı MP3'e dönüştürür."""
    return convert_audio(input_path, output_folder, "mp3")

def convert_mp4_to_avi(input_path, output_folder):
    return convert_video(input_path, output_folder, 'avi')

# --- YouTube İndirme İşleyici ---
def handle_youtube_download(form_data, output_folder):
    """YouTube URL'sini alır, sesi indirir ve dosyayı döndürür."""
    youtube_url = form_data.get('youtube_url')
    if not youtube_url:
        raise ValueError("YouTube URL'si sağlanmadı.")

    ydl_opts = {
        'format': 'bestaudio/best',
        'postprocessors': [{'key': 'FFmpegExtractAudio', 'preferredcodec': 'mp3', 'preferredquality': '192'}],
        'outtmpl': os.path.join(output_folder, '%(title)s.%(ext)s'),
        'ffmpeg_location': ffmpeg_path
    }

    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(youtube_url, download=True)
            filename = ydl.prepare_filename(info)
            # Uzantıyı .mp3 olarak değiştir
            base, _ = os.path.splitext(filename)
            output_path = base + '.mp3'
            
            if os.path.exists(output_path):
                 logging.info(f"YouTube ses indirme başarılı: {output_path}")
                 return output_path
            # Bazen dosya adı farklı olabilir, klasördeki mp3'ü bul
            for f in os.listdir(output_folder):
                if f.endswith('.mp3'):
                    return os.path.join(output_folder, f)
            raise FileNotFoundError("İndirilen MP3 dosyası bulunamadı.")
            
    except Exception as e:
        logging.error(f"YouTube'dan ses indirme hatası: {e}")
        return None

# --- DESTEKLENEN DÖNÜŞÜMLER ---
# Bu sözlük yapısı, yeni dönüştürücülerin kolayca eklenmesini sağlar.
# Her anahtar, formda görünecek olan 'conversion_type' değeridir.
CONVERTERS = {
    'word-to-pdf': {
        'display_name': "Word'den PDF'e (.docx → .pdf)",
        'function': convert_word_to_pdf,
        'allowed_extensions': {'docx'},
        'output_format': 'pdf'
    },
    'pdf-to-word': {
        'display_name': "PDF'ten Word'e (.pdf → .docx)",
        'function': convert_pdf_to_word,
        'allowed_extensions': {'pdf'},
        'output_format': 'docx'
    },
    'word-to-txt': {
        'display_name': "Word'den Metine (.docx → .txt)",
        'function': convert_word_to_txt,
        'allowed_extensions': {'docx'},
        'output_format': 'txt'
    },
    'txt-to-word': {
        'display_name': "Metinden Word'e (.txt → .docx)",
        'function': convert_txt_to_word,
        'allowed_extensions': {'txt'},
        'output_format': 'docx'
    },
    'pdf-to-jpg': {
        'display_name': "PDF'ten JPG'ye (.pdf → .jpg/.zip)",
        'function': convert_pdf_to_jpg,
        'allowed_extensions': {'pdf'},
        'output_format': 'zip' # Çoklu sayfalar için ZIP dönebilir
    },
    'jpg-to-pdf': {
        'display_name': "JPG'den PDF'e (.jpg → .pdf)",
        'function': convert_jpg_to_pdf,
        'allowed_extensions': {'jpg', 'jpeg'},
        'output_format': 'pdf'
    },
    'jpg-to-png': {
        'display_name': "JPG'den PNG'ye (.jpg → .png)",
        'function': convert_jpg_to_png,
        'allowed_extensions': {'jpg', 'jpeg'},
        'output_format': 'png'
    },
    'png-to-jpg': {
        'display_name': "PNG'den JPG'ye (.png → .jpg)",
        'function': convert_png_to_jpg,
        'allowed_extensions': {'png'},
        'output_format': 'jpg'
    },
    'wav-to-mp3': {
        'display_name': "WAV'dan MP3'e (.wav → .mp3)",
        'function': convert_wav_to_mp3,
        'allowed_extensions': {'wav'},
        'output_format': 'mp3'
    },
    'mp4-to-avi': {
        'display_name': "MP4'ten AVI'ye (.mp4 → .avi)",
        'function': convert_mp4_to_avi,
        'allowed_extensions': {'mp4'},
        'output_format': 'avi'
    },
    'json-to-xml': {
        'display_name': "JSON'dan XML'e (.json → .xml)",
        'function': convert_json_to_xml,
        'allowed_extensions': {'json'},
        'output_format': 'xml'
    },
    'xml-to-json': {
        'display_name': "XML'den JSON'a (.xml → .json)",
        'function': convert_xml_to_json,
        'allowed_extensions': {'xml'},
        'output_format': 'json'
    },
    'rar-to-zip': {
        'display_name': "RAR'dan ZIP'e",
        'allowed_extensions': ["rar"],
        'function': convert_rar_to_zip,
        'output_format': 'zip'
    },
    # --- Çevrimiçi Medya İndiricileri ---
    "spotify-downloader": {
        "display_name": "Spotify'dan MP3 İndir",
        "is_online_service": True,
        # Bu fonksiyon doğrudan route üzerinden yönetiliyor, burada bir işlem yapmasına gerek yok.
        "function": None, 
        "form_fields": [
            {"name": "spotify-links", "label": "Spotify Şarkı/Playlist Linkleri (her satıra bir tane)"}
        ]
    },
    "youtube-audio-downloader": {
        "display_name": "YouTube'dan Ses İndir",
        "is_online_service": True,
        "function": handle_youtube_download,
        "form_fields": [
            {"name": "youtube_url", "label": "YouTube Video URL"}
        ]
    }
}

# Sadece Windows'ta çalışacak dönüştürücüleri ekle
if platform.system() == "Windows":
    try:
        import win32com.client
        logging.info("Windows algılandı. Office dönüştürücüleri ekleniyor.")
        
        CONVERTERS['excel-to-pdf'] = {
            'display_name': "Excel'den PDF'e (.xlsx → .pdf)",
            'function': convert_excel_to_pdf,
            'allowed_extensions': {'xlsx'},
            'output_format': 'pdf'
        }
        CONVERTERS['powerpoint-to-pdf'] = {
            'display_name': "PowerPoint'ten PDF'e (.pptx → .pdf)",
            'function': convert_powerpoint_to_pdf,
            'allowed_extensions': {'pptx', 'ppt'},
            'output_format': 'pdf'
        }
    except ImportError:
        logging.warning("Windows algılandı ancak 'pypiwin32' kütüphanesi bulunamadı. Office dönüştürücüleri devre dışı bırakıldı.")


# --- Spotify İndirme Durum Takibi ---
# Her session_id için ayrı bir durum ve dosya listesi tutulur
# Spotify sessions artık SessionManager tarafından yönetiliyor

def get_spotify_track_info(track_url):
    """Spotify track URL'sinden şarkı adı ve sanatçıyı çeker."""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    try:
        response = requests.get(track_url, headers=headers, timeout=10)
        if response.status_code != 200:
            return None, None
        soup = BeautifulSoup(response.content, 'html.parser')
        title = soup.find('meta', property='og:title')
        description = soup.find('meta', property='og:description')
        
        song_name = title['content'] if title else 'Bilinmeyen Şarkı'
        
        artist = 'Bilinmeyen Sanatçı'
        if description:
            # "Listen to X on Spotify. Y · Song · 2023." formatından sanatçıyı al
            desc_content = description['content']
            parts = desc_content.split('·')
            if len(parts) > 1:
                artist_candidate = parts[0].replace('Listen to ', '').replace(f' on Spotify', '').strip()
                # Eğer sanatçı adı çok uzunsa, muhtemelen şarkı adıdır, ilk bölümü al.
                if len(artist_candidate) > len(song_name) + 5:
                    artist = song_name
                else:
                    artist = artist_candidate

        # Bazen başlık "Şarkı - Sanatçı" formatında gelir
        if ' - ' in song_name and artist == 'Bilinmeyen Sanatçı':
            parts = song_name.split(' - ', 1)
            song_name = parts[0]
            artist = parts[1]

        return song_name, artist
    except Exception as e:
        logging.error(f"Spotify bilgisi alınamadı ({track_url}): {e}")
        return None, None

def download_youtube_audio(search_query, output_path, song_name, session_id):
    """Verilen arama sorgusu ile YouTube'dan en iyi ses sonucunu indirir."""
    session = session_manager.get_session(session_id)
    if not session: return

    def progress_hook(d):
        if d['status'] == 'downloading':
            percent = d.get('_percent_str', '0%').strip().replace('%', '')
            session['status'][song_name] = f"İndiriliyor... {percent}%"
        elif d['status'] == 'finished':
            session['status'][song_name] = "İşleniyor..."

    try:
        session['status'][song_name] = "YouTube'da aranıyor..."
        ydl_opts = {
            'format': 'bestaudio/best',
            'progress_hooks': [progress_hook],
            'postprocessors': [{'key': 'FFmpegExtractAudio', 'preferredcodec': 'mp3', 'preferredquality': '192'}],
            'outtmpl': os.path.join(output_path, f'{search_query}.%(ext)s'),
            'default_search': 'ytsearch5',
            'force_ipv4': True,  # IPv4 kullanmaya zorla
            'verbose': True, # Hata ayıklama için ayrıntılı çıktı
            'ffmpeg_location': ffmpeg_path # FFmpeg yolunu burada belirt
        }
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            # Arama sonucundan bilgi al, indirme yapma
            info = ydl.extract_info(search_query, download=False)
            
            # Gerçek indirme için seçenekleri ayarla (sessiz modda)
            ydl.params['quiet'] = True
            ydl.params['verbose'] = False
            ydl.params['no_warnings'] = True
            
            # Videoyu indir
            ydl.download([info['entries'][0]['id']])

            # İndirme sonrası dosya adını bul ve yeniden adlandır
            original_ext = info['entries'][0]['ext']
            downloaded_file = ydl.prepare_filename(info['entries'][0]).rsplit('.', 1)[0] + '.mp3'

            # Dosya adını güvenli hale getir
            safe_filename = secure_filename(f"{search_query}.mp3")
            final_path = os.path.join(output_path, safe_filename)
            
            # Eğer dosya zaten varsa ismini değiştirerek kaydet
            if os.path.exists(final_path):
                 final_path = os.path.join(output_path, f"{datetime.now().strftime('%H%M%S')}_{safe_filename}")

            if os.path.exists(downloaded_file):
                os.rename(downloaded_file, final_path)
            
            if os.path.exists(final_path):
                session['status'][song_name] = "Tamamlandı"
                session['files'].append(final_path)
                logging.info(f"'{search_query}' başarıyla indirildi: {final_path}")
            else:
                 raise FileNotFoundError(f"İndirilen dosya bulunamadı: {downloaded_file}")

    except yt_dlp.utils.DownloadError as e:
        logging.error(f"yt-dlp indirme hatası ({search_query}): {e}")
        # Hata mesajını daha anlaşılır hale getir
        error_message = str(e)
        if 'HTTP Error 403' in error_message:
            error_message = "YouTube erişimi engelledi (403 Forbidden)."
        elif 'HTTP Error 410' in error_message:
            error_message = "Video artık mevcut değil (410 Gone)."
        elif 'unable to download video data' in error_message:
            error_message = "Video verisi indirilemiyor. Video özel veya silinmiş olabilir."
        else:
            # Hatanın başını al, çok uzun olmasın
            error_message = re.sub(r'\[[^\]]+\]', '', error_message).strip().split('\n')[-1]

        session['status'][song_name] = f"Hata: {error_message[:100]}"
    except Exception as e:
        logging.error(f"Genel YouTube indirme hatası ({search_query}): {e}")
        session['status'][song_name] = f"Hata: {str(e)[:100]}..."


@app.route('/', methods=['GET', 'POST'])
def index():
    """Ana sayfa. Dosya yükleme formunu gösterir ve dönüştürme isteğini işler."""
    # İşletim sistemi ve FFmpeg durumunu şablona gönder
    is_windows = platform.system() == "Windows"
    ffmpeg_available = bool(ffmpeg_path)

    if request.method == 'POST':
        # Disk alanını kontrol et
        try:
            check_disk_space()
        except Exception as e:
            flash(str(e), 'error')
            return redirect(url_for('index'))
        conversion_type = request.form.get('conversion_type')
        if not conversion_type or conversion_type not in CONVERTERS:
            flash('Geçersiz dönüştürme türü seçtiniz.', 'error')
            return redirect(request.referrer or url_for('index'))
        
        converter_info = CONVERTERS[conversion_type]

        try:
            # Her işlem için benzersiz bir klasör oluştur
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            # Güvenli bir temel ad oluştur
            base_name = re.sub(r'[^a-zA-Z0-9_.-]', '', f"job_{conversion_type}")
            job_folder = os.path.join(app.config['DOWNLOAD_FOLDER'], f"{timestamp}_{base_name}")
            os.makedirs(job_folder, exist_ok=True)
            
            output_path = None

            # Çevrimiçi servisler dosya yüklemesi gerektirmez
            if converter_info.get('is_online_service'):
                 if conversion_type == 'youtube-audio-downloader':
                    output_path = converter_info['function'](request.form, job_folder)
                 else:
                    # Spotify gibi diğer online servisler kendi rotaları üzerinden yönetilir.
                    # Bu POST isteği buraya gelmemeli.
                    flash("Beklenmeyen bir istek yapıldı.", 'error')
                    return redirect(url_for('index'))
            else: # Dosya tabanlı dönüştürücüler
                uploaded_file = request.files.get('file')
                if not uploaded_file or uploaded_file.filename == '':
                    flash('Lütfen bir dosya seçin.', 'error')
                    return redirect(request.referrer or url_for('index'))

                original_filename = secure_filename(uploaded_file.filename)
                file_extension = '.' in original_filename and original_filename.rsplit('.', 1)[1].lower()

                if file_extension not in converter_info['allowed_extensions']:
                    allowed = ", ".join(converter_info['allowed_extensions'])
                    flash(f"Hatalı dosya türü. Lütfen bir {allowed} dosyası yükleyin.", 'error')
                    return redirect(request.referrer or url_for('index'))
                
                input_path = os.path.join(job_folder, original_filename)
                uploaded_file.save(input_path)
                logging.info(f"Dosya geçici olarak '{input_path}' konumuna kaydedildi.")
                output_path = converter_info['function'](input_path, job_folder)

            # Sonucu kullanıcıya gönder
            if output_path and os.path.exists(output_path):
                logging.info(f"Dönüştürülen dosya '{output_path}' indirilmek üzere gönderiliyor.")
                return send_file(output_path, as_attachment=True)
            else:
                flash("Dosya dönüştürme sırasında bir hata oluştu veya dönüştürücü bir dosya döndürmedi. Lütfen tekrar deneyin.", 'error')
                return redirect(request.referrer or url_for('index'))

        except Exception as e:
            logging.error(f"İşlem sırasında beklenmedik bir hata oluştu: {e}")
            import traceback
            logging.error(traceback.format_exc())
            
            error_message = "Beklenmedik bir sunucu hatası oluştu. Lütfen yönetici ile iletişime geçin."
            if "unrar' programı sisteminizde bulunamadı" in str(e):
                error_message = "RAR dönüştürme başarısız: 'unrar' programı sistemde kurulu veya erişilebilir değil."
            elif isinstance(e, ValueError):
                error_message = str(e)
            
            flash(error_message, 'error')
            return redirect(request.referrer or url_for('index'))

    # GET isteği için her zaman sayfayı render et
    return render_template('index.html', 
                           converters=CONVERTERS, 
                           is_windows=is_windows, 
                           ffmpeg_available=ffmpeg_available)

@app.route('/download_spotify', methods=['POST'])
def download_spotify_route():
    """Spotify linklerini alır ve indirme işlemini başlatır."""
    # Rate limiting kontrol
    if limiter:
        try:
            limiter.test_request()
        except Exception:
            return jsonify({'error': 'Çok fazla istek gönderdiniz. Lütfen bekleyin.'}), 429
    
    # Disk alanını kontrol et
    try:
        check_disk_space()
    except Exception as e:
        return jsonify({'error': str(e)}), 507
    
    data = request.get_json()
    links_text = data.get('links', '')
    if not links_text:
        return jsonify({'error': 'Link metni boş olamaz.'}), 400

    track_urls = re.findall(r'https://open\.spotify\.com/track/[a-zA-Z0-9]+', links_text)
    if not track_urls:
        return jsonify({'error': 'Geçerli Spotify şarkı linki bulunamadı.'}), 400
    
    # Çok fazla şarkı kontrolü
    if len(track_urls) > 20:
        return jsonify({'error': 'Maksimum 20 şarkı aynı anda indirilebilir.'}), 400
    
    session_id = f"spotify_{datetime.now().strftime('%Y%m%d%H%M%S')}_{os.urandom(4).hex()}"
    session_folder = os.path.join(app.config['DOWNLOAD_FOLDER'], session_id)
    os.makedirs(session_folder)

    session_manager.create_session(session_id, {'status': {}, 'files': [], 'is_complete': False})
    
    # Thread pool executor kullanarak indirme işlemini başlat
    future = executor.submit(spotify_download_thread, track_urls, session_folder, session_id)
    
    return jsonify({'message': 'İndirme başlatıldı.', 'session_id': session_id})

def spotify_download_thread(track_urls, session_folder, session_id):
    """Arka planda Spotify şarkılarını indiren thread fonksiyonu."""
    session = session_manager.get_session(session_id)
    if not session:
        logging.error(f"Session bulunamadı: {session_id}")
        return
        
    try:
        for url in track_urls:
            # Session'ın hala var olduğunu kontrol et
            session = session_manager.get_session(session_id)
            if not session:
                logging.warning(f"Session süresi doldu: {session_id}")
                break
                
            song_name, artist = get_spotify_track_info(url)
            if song_name and artist:
                display_name = f"{artist} - {song_name}"
                session['status'][display_name] = "Sırada"
                session_manager.update_session(session_id, session)
                
                search_query = f"{artist} {song_name}"
                download_youtube_audio(search_query, session_folder, display_name, session_id)
            else:
                session['status'][url] = "Hata: Şarkı bilgisi alınamadı"
                session_manager.update_session(session_id, session)
            
            time.sleep(1)  # Rate limiting
        
        # Session'ı tamamlandı olarak işaretle
        session = session_manager.get_session(session_id)
        if session:
            session['is_complete'] = True
            session_manager.update_session(session_id, session)
            logging.info(f"Spotify indirme oturumu ({session_id}) tamamlandı.")
            
    except Exception as e:
        logging.error(f"Spotify indirme thread hatası ({session_id}): {e}")
        session = session_manager.get_session(session_id)
        if session:
            session['is_complete'] = True
            session['error'] = str(e)
            session_manager.update_session(session_id, session)


@app.route('/spotify_status/<session_id>')
def spotify_status_route(session_id):
    """Belirli bir Spotify indirme oturumunun durumunu döndürür."""
    session = session_manager.get_session(session_id)
    if not session:
        return jsonify({'error': 'Oturum bulunamadı veya süresi doldu.'}), 404
    
    zip_ready = session['is_complete'] and any(f.endswith('.mp3') for f in session.get('files', []))
    
    return jsonify({
        'status': session.get('status', {}),
        'is_complete': session.get('is_complete', False),
        'zip_ready': zip_ready,
        'error': session.get('error')
    })

@app.route('/download_spotify_zip/<session_id>')
def download_spotify_zip_route(session_id):
    """Tamamlanan Spotify indirmelerini bir ZIP dosyası olarak sunar."""
    session = session_manager.get_session(session_id)
    if not session or not session.get('is_complete'):
        return "Oturum bulunamadı, süresi doldu veya henüz tamamlanmadı.", 404

    downloaded_files = session.get('files', [])
    if not downloaded_files:
        return "İndirilecek dosya bulunamadı.", 404

    zip_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f"{session_id}.zip")
    try:
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for f_path in downloaded_files:
                if os.path.exists(f_path):
                    zipf.write(f_path, os.path.basename(f_path))

        return send_file(zip_path, as_attachment=True)
    except Exception as e:
        logging.error(f"ZIP oluşturma hatası ({session_id}): {e}")
        return "ZIP dosyası oluşturulamadı.", 500


# --- ADMİN VE MONİTORİNG ---

@app.route('/admin/status')
def admin_status():
    """Sistem durumu endpoint'i (admin için)"""
    try:
        stats = get_system_stats()
        
        # Downloads klasöründeki toplam dosya sayısı
        total_files = 0
        total_folders = 0
        for root, dirs, files in os.walk(app.config['DOWNLOAD_FOLDER']):
            total_files += len(files)
            total_folders += len(dirs)
        
        return jsonify({
            'status': 'healthy',
            'system': stats,
            'downloads': {
                'total_files': total_files,
                'total_folders': total_folders,
                'retention_hours': app.config['FILE_RETENTION_HOURS']
            },
            'config': {
                'max_concurrent_downloads': app.config['MAX_CONCURRENT_DOWNLOADS'],
                'cleanup_interval_hours': app.config['CLEANUP_INTERVAL_HOURS'],
                'disk_warning_percent': app.config['DISK_USAGE_WARNING_PERCENT'],
                'disk_critical_percent': app.config['DISK_USAGE_CRITICAL_PERCENT']
            },
            'timestamp': datetime.now().isoformat()
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }), 500

@app.route('/admin/cleanup', methods=['POST'])
def admin_cleanup():
    """Manuel dosya temizleme endpoint'i (admin için)"""
    try:
        cleaned_count = cleanup_old_files()
        session_manager.cleanup_expired_sessions()
        
        return jsonify({
            'status': 'success',
            'cleaned_files': cleaned_count,
            'message': f'{cleaned_count} eski dosya/klasör temizlendi',
            'timestamp': datetime.now().isoformat()
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }), 500

@app.errorhandler(413)
def request_entity_too_large(error):
    """Dosya boyutu limitini aşan istekler için hata."""
    flash(f"Yüklemeye çalıştığınız dosya çok büyük. Maksimum boyut: {app.config['MAX_CONTENT_LENGTH'] // 1024 // 1024} MB.", 'error')
    return redirect(url_for('index'))


if __name__ == '__main__':
    # Geliştirme ortamı için debug modunu aç.
    # Production ortamında bir WSGI sunucusu (Gunicorn, Waitress vb.) kullanılmalıdır.
    app.run(debug=True) 